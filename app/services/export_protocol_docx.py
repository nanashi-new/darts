from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.services.export_protocol_xlsx import ProtocolData


def export_protocol_docx(path: str, data: ProtocolData) -> None:
    """Export a formatted protocol to DOCX."""
    doc = Document()

    # Logo
    if data.logo_path and os.path.isfile(data.logo_path):
        try:
            doc.add_picture(data.logo_path, width=Cm(3))
        except Exception:  # noqa: BLE001
            pass

    # Org name lines
    org_lines = data.org_name.split("\n") if data.org_name else []
    for line in org_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line.strip())
        run.bold = True
        run.font.size = Pt(12)

    # City
    if data.city:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(data.city)

    # Competition title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(data.competition_title)
    run.bold = True
    run.font.size = Pt(14)

    # PROTOKOL REZULTATOV
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u041f\u0420\u041e\u0422\u041e\u041a\u041e\u041b \u0420\u0415\u0417\u0423\u041b\u042c\u0422\u0410\u0422\u041e\u0412")
    run.bold = True
    run.font.size = Pt(14)

    # Category
    if data.category:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(data.category)

    # Format type
    format_labels = {
        "classification": "\u041a\u043b\u0430\u0441\u0441\u0438\u0444\u0438\u043a\u0430\u0446\u0438\u044f",
        "501": "501 - \u043e\u0434\u0438\u043d\u043e\u0447\u043d\u044b\u0439 \u0440\u0430\u0437\u0440\u044f\u0434",
        "norms": "\u0421\u0434\u0430\u0447\u0430 \u043d\u043e\u0440\u043c\u0430\u0442\u0438\u0432\u043e\u0432",
    }
    format_label = format_labels.get(data.format_type, data.format_type)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(format_label)

    # Venue + date
    venue_date = f"{data.venue}, {data.date}" if data.venue else data.date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(venue_date)

    # Jury section
    if data.jury:
        p = doc.add_paragraph()
        run = p.add_run("\u0421\u0443\u0434\u0435\u0439\u0441\u043a\u0430\u044f \u043a\u043e\u043b\u043b\u0435\u0433\u0438\u044f:")
        run.bold = True

        jury_table = doc.add_table(rows=len(data.jury), cols=4)
        for row_idx, jury_member in enumerate(data.jury):
            jury_table.cell(row_idx, 0).text = jury_member.get("position", "")
            jury_table.cell(row_idx, 1).text = jury_member.get("name", "")
            jury_table.cell(row_idx, 2).text = jury_member.get("category", "")
            jury_table.cell(row_idx, 3).text = jury_member.get("city", "")

    # Results table
    if data.format_type == "501":
        columns = [
            "\u041c\u0435\u0441\u0442\u043e",
            "\u0424\u0430\u043c\u0438\u043b\u0438\u044f, \u0418\u043c\u044f",
            "\u0413\u043e\u0434 \u0440\u043e\u0436\u0434\u0435\u043d\u0438\u044f",
            "\u0417\u0432\u0430\u043d\u0438\u0435, \u0440\u0430\u0437\u0440\u044f\u0434",
            "\u0421\u0443\u0431\u044a\u0435\u043a\u0442 \u0420\u0424, \u0433\u043e\u0440\u043e\u0434",
            "\u0422\u0440\u0435\u043d\u0435\u0440",
            "\u0412\u044b\u043f\u043e\u043b\u043d\u0435\u043d \u0440\u0430\u0437\u0440\u044f\u0434",
        ]
        result_keys = ["place", "fio", "birth_year", "current_rank", "region", "coach", "rank_achieved"]
    else:
        columns = [
            "\u041c\u0435\u0441\u0442\u043e",
            "\u0424\u0418\u041e",
            "\u0413/\u0420",
            "\u0422\u0440\u0435\u043d\u0435\u0440",
            "\u041d\u0430\u0431\u043e\u0440 \u043e\u0447\u043a\u043e\u0432",
            "\u0421\u0435\u043a\u0442\u043e\u0440 20",
            "\u0411\u043e\u043b\u044c\u0448\u043e\u0439 \u0440\u0430\u0443\u043d\u0434",
            "\u0418\u0442\u043e\u0433\u043e",
            "\u0412\u044b\u043f\u043e\u043b\u043d\u0435\u043d\u043d\u044b\u0439 \u0440\u0430\u0437\u0440\u044f\u0434",
        ]
        result_keys = [
            "place", "fio", "birth_year", "coach",
            "score_set", "score_sector20", "score_big_round",
            "points_total", "rank_achieved",
        ]

    num_results = len(data.results)
    results_table = doc.add_table(rows=num_results + 1, cols=len(columns))
    results_table.style = "Table Grid"

    # Header row
    for col_idx, col_name in enumerate(columns):
        cell = results_table.cell(0, col_idx)
        cell.text = col_name
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    for row_idx, result in enumerate(data.results, start=1):
        for col_idx, key in enumerate(result_keys):
            value = result.get(key, "")
            results_table.cell(row_idx, col_idx).text = str(value) if value is not None else ""

    # Signature section
    doc.add_paragraph()  # blank

    chief_judge_name = ""
    chief_secretary_name = ""
    for jury_member in data.jury:
        pos = (jury_member.get("position") or "").lower()
        if "\u0433\u043b\u0430\u0432\u043d\u044b\u0439 \u0441\u0443\u0434\u044c\u044f" in pos or "chief judge" in pos:
            chief_judge_name = jury_member.get("name", "")
        elif "\u0433\u043b\u0430\u0432\u043d\u044b\u0439 \u0441\u0435\u043a\u0440\u0435\u0442\u0430\u0440\u044c" in pos or "secretary" in pos:
            chief_secretary_name = jury_member.get("name", "")

    doc.add_paragraph(f"\u0413\u043b\u0430\u0432\u043d\u044b\u0439 \u0441\u0443\u0434\u044c\u044f _________ {chief_judge_name}")
    doc.add_paragraph(f"\u0413\u043b\u0430\u0432\u043d\u044b\u0439 \u0441\u0435\u043a\u0440\u0435\u0442\u0430\u0440\u044c _________ {chief_secretary_name}")

    doc.save(path)
