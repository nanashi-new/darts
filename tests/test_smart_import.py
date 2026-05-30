from __future__ import annotations

import pytest
from docx import Document

from app.services.import_protocol_docx import parse_tables_from_docx

pytestmark = pytest.mark.integration


def _add_classification_table(doc: Document, rows_data: list[list[str]]) -> None:
    """Add a classification-style table (16 columns) to a document."""
    headers = [
        "место", "ФИО", "", "", "Г/Р",
        "тренер", "", "", "набор очков", "",
        "сектор 20", "", "Большой раунд", "", "итого",
        "выполненный разряд",
    ]
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows_data), cols=num_cols)
    # Set header row
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    # Set data rows
    for row_idx, row in enumerate(rows_data, start=1):
        for col_idx, val in enumerate(row):
            if col_idx < num_cols:
                table.rows[row_idx].cells[col_idx].text = val


def _add_501_table(doc: Document, rows_data: list[list[str]]) -> None:
    """Add a 501-style table (7 columns) to a document."""
    headers = [
        "МЕСТО", "Фамилия, Имя", "Год рождения",
        "Звание, разряд", "Субъект РФ, город", "тренер", "Выполнен разряд",
    ]
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows_data), cols=num_cols)
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row_idx, row in enumerate(rows_data, start=1):
        for col_idx, val in enumerate(row):
            if col_idx < num_cols:
                table.rows[row_idx].cells[col_idx].text = val


def _add_jury_table(doc: Document) -> None:
    """Add a 4-column jury table to a document."""
    headers = ["должность", "ФИО", "город", "формат"]
    table = doc.add_table(rows=2, cols=4)
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = "Главный судья"
    table.rows[1].cells[1].text = "Иванов И.И."
    table.rows[1].cells[2].text = "Москва"
    table.rows[1].cells[3].text = "501"


def test_parse_docx_classification_table(tmp_path) -> None:
    """Classification table with 2 data rows should produce 1 TableBlock."""
    doc = Document()
    doc.add_paragraph("Юниоры 15-17 лет")

    rows_data = [
        ["1", "Иванов Иван", "", "", "2008", "Петров П.П.", "", "", "150", "", "80", "", "60", "", "290", "1р"],
        ["2", "Сидоров Сидор", "", "", "2009", "Козлов К.К.", "", "", "130", "", "70", "", "50", "", "250", "2р"],
    ]
    _add_classification_table(doc, rows_data)

    path = tmp_path / "test_classification.docx"
    doc.save(str(path))

    result = parse_tables_from_docx(str(path))

    assert len(result) == 1
    block = result[0]
    assert block.sheet_name == "Юниоры 15-17 лет"
    assert len(block.rows) == 2
    assert block.rows[0]["fio"] == "Иванов Иван"
    assert block.rows[0]["birth"] == "2008"
    assert block.rows[0]["coach"] == "Петров П.П."
    assert block.rows[0]["place"] == "1"
    assert block.rows[0]["score_set"] == "150"
    assert block.rows[0]["score_sector20"] == "80"
    assert block.rows[0]["score_big_round"] == "60"
    assert block.confidence == 1.0
    assert block.needs_mapping is False


def test_parse_docx_skips_jury_table(tmp_path) -> None:
    """Jury tables (4 cols with jury keywords) should be skipped."""
    doc = Document()
    doc.add_paragraph("Мужчины 18+ лет")

    _add_jury_table(doc)

    rows_data = [
        ["1", "Иванов Иван", "1990", "КМС", "Москва", "Тренеров Т.Т.", "КМС"],
    ]
    _add_501_table(doc, rows_data)

    path = tmp_path / "test_jury_skip.docx"
    doc.save(str(path))

    result = parse_tables_from_docx(str(path))

    assert len(result) == 1
    block = result[0]
    assert block.rows[0]["fio"] == "Иванов Иван"
    assert block.rows[0]["birth"] == "1990"
    assert block.rows[0]["coach"] == "Тренеров Т.Т."
    assert block.rows[0]["place"] == "1"


def test_parse_docx_empty_file(tmp_path) -> None:
    """An empty DOCX file should return empty list without raising."""
    doc = Document()
    path = tmp_path / "empty.docx"
    doc.save(str(path))

    result = parse_tables_from_docx(str(path))
    assert result == []


def test_parse_docx_nonexistent_file() -> None:
    """A non-existent file should return empty list without raising."""
    result = parse_tables_from_docx("/nonexistent/path/file.docx")
    assert result == []


def test_parse_docx_501_table(tmp_path) -> None:
    """501-format table should be parsed with fio, birth, coach, place."""
    doc = Document()
    doc.add_paragraph("Женщины 18+ лет")

    rows_data = [
        ["1", "Петрова Мария", "1995", "1р", "СПб", "Сидоров С.С.", "КМС"],
        ["2", "Козлова Анна", "1998", "2р", "Москва", "Иванов И.И.", "1р"],
    ]
    _add_501_table(doc, rows_data)

    path = tmp_path / "test_501.docx"
    doc.save(str(path))

    result = parse_tables_from_docx(str(path))

    assert len(result) == 1
    block = result[0]
    assert block.sheet_name == "Женщины 18+ лет"
    assert len(block.rows) == 2
    assert block.rows[0]["fio"] == "Петрова Мария"
    assert block.rows[0]["birth"] == "1995"
    assert block.rows[0]["coach"] == "Сидоров С.С."
    assert block.rows[1]["fio"] == "Козлова Анна"
    # 501 tables don't have score columns
    assert block.rows[0]["score_set"] is None


def test_parse_docx_multiple_categories(tmp_path) -> None:
    """Multiple categories produce separate TableBlocks."""
    doc = Document()

    doc.add_paragraph("Юниоры 15-17 лет")
    _add_501_table(doc, [["1", "Иванов Иван", "2008", "", "", "Тренер А", ""]])

    doc.add_paragraph("Юниорки 15-17 лет")
    _add_501_table(doc, [["1", "Петрова Мария", "2009", "", "", "Тренер Б", ""]])

    path = tmp_path / "test_multi.docx"
    doc.save(str(path))

    result = parse_tables_from_docx(str(path))

    assert len(result) == 2
    assert result[0].sheet_name == "Юниоры 15-17 лет"
    assert result[1].sheet_name == "Юниорки 15-17 лет"


# --- PDF parser tests ---

from app.services.import_protocol_pdf import parse_lines, parse_tables_from_pdf
from app.services.import_pipeline import detect_format


def test_parse_pdf_basic() -> None:
    """parse_lines extracts rows from structured text with place, FIO, birth, scores."""
    text = (
        "1 Иванов Иван Иванович 2008 150 80 60 290 1р\n"
        "2 Сидоров Сидор 2009 130 70 50 250 2р\n"
    )
    rows, warnings, _cat = parse_lines(text)
    assert len(rows) == 2
    assert rows[0]["fio"] == "Иванов Иван Иванович"
    assert rows[0]["place"] == "1"
    assert rows[0]["birth"] == "2008"
    assert rows[0]["score_set"] == "150"
    assert rows[0]["score_sector20"] == "80"
    assert rows[0]["score_big_round"] == "60"
    assert rows[1]["fio"] == "Сидоров Сидор"
    assert rows[1]["place"] == "2"


def test_parse_pdf_category_detection() -> None:
    """Category keywords like 'мужчины'/'женщины' are detected from text."""
    text = (
        "мужчины 18+ лет\n"
        "1 Иванов Иван 1990 200 100 80 380 КМС\n"
        "женщины 18+ лет\n"
        "1 Петрова Мария 1995 180 90 70 340 1р\n"
    )
    rows, warnings, last_cat = parse_lines(text)
    # parse_lines returns the last category detected
    assert last_cat == "женщины 18+ лет"
    # Both rows parsed (categories don't split rows in parse_lines)
    assert len(rows) == 2


def test_parse_pdf_empty_file(tmp_path) -> None:
    """Non-existent file returns empty list without raising."""
    result = parse_tables_from_pdf("/nonexistent/path/file.pdf")
    assert result == []

    # Empty file (not a valid PDF)
    empty_file = tmp_path / "empty.pdf"
    empty_file.write_bytes(b"")
    result = parse_tables_from_pdf(str(empty_file))
    assert result == []


def test_detect_format_docx() -> None:
    """detect_format returns 'docx' for .docx files."""
    assert detect_format("file.docx") == "docx"
    assert detect_format("path/to/File.DOCX") == "docx"


def test_detect_format_pdf() -> None:
    """detect_format returns 'pdf' for .pdf files."""
    assert detect_format("file.pdf") == "pdf"
    assert detect_format("path/to/File.PDF") == "pdf"


def test_detect_format_preserves_existing() -> None:
    """detect_format still correctly handles csv, json, xlsx."""
    assert detect_format("f.csv") == "csv"
    assert detect_format("f.json") == "json"
    assert detect_format("f.xlsx") == "xlsx"


# --- Import modes tests (require in-memory DB) ---

import sqlite3

from app.db.schema import initialize_schema
from app.services.import_modes import (
    import_multi_tournament,
    import_players_only,
    import_update_players,
)
from app.services.import_xlsx import TableBlock


def _make_db() -> sqlite3.Connection:
    """Create an in-memory database with full schema."""
    conn = sqlite3.connect(":memory:")
    conn.row_factory = sqlite3.Row
    initialize_schema(conn)
    return conn


def _make_block(
    rows: list[dict[str, object]],
    sheet_name: str = "test",
) -> TableBlock:
    """Create a minimal TableBlock with given rows."""
    return TableBlock(
        sheet_name=sheet_name,
        start_row=1,
        end_row=len(rows),
        header_mapping={},
        rows=rows,
        warnings=[],
        errors=[],
        needs_mapping=False,
        confidence=1.0,
        missing_required_columns=[],
    )


@pytest.mark.integration
def test_import_players_only_creates_new() -> None:
    """import_players_only creates new players and reports counts."""
    conn = _make_db()
    block = _make_block([
        {"fio": "Иванов Иван", "birth": "2010", "coach": "Тренер А"},
        {"fio": "Петров Петр", "birth": "2011", "coach": "Тренер Б"},
    ])

    report = import_players_only(connection=conn, blocks=[block])

    assert report.created == 2
    assert report.existing == 0


@pytest.mark.integration
def test_import_players_only_detects_existing() -> None:
    """import_players_only detects an existing player and only creates new ones."""
    conn = _make_db()
    # Pre-create a player
    conn.execute(
        "INSERT INTO players (last_name, first_name, middle_name, birth_date) "
        "VALUES (?, ?, ?, ?)",
        ("иванов", "иван", None, "2010"),
    )
    conn.commit()

    block = _make_block([
        {"fio": "Иванов Иван", "birth": "2010", "coach": "Тренер А"},
        {"fio": "Петров Петр", "birth": "2011", "coach": "Тренер Б"},
    ])

    report = import_players_only(connection=conn, blocks=[block])

    assert report.created == 1
    assert report.existing == 1


@pytest.mark.integration
def test_import_update_players_fills_empty_coach() -> None:
    """import_update_players fills empty coach field for an existing player."""
    conn = _make_db()
    # Pre-create a player with no coach
    conn.execute(
        "INSERT INTO players (last_name, first_name, middle_name, birth_date, coach) "
        "VALUES (?, ?, ?, ?, ?)",
        ("иванов", "иван", None, "2010", None),
    )
    conn.commit()

    block = _make_block([
        {"fio": "Иванов Иван", "birth": "2010", "coach": "Тренеров Т.Т."},
    ])

    report = import_update_players(connection=conn, blocks=[block])

    assert report.updated == 1
    # Verify the coach was actually written to DB
    row = conn.execute("SELECT coach FROM players WHERE last_name = 'иванов'").fetchone()
    assert row["coach"] == "Тренеров Т.Т."


@pytest.mark.integration
def test_import_update_players_does_not_overwrite_existing_coach() -> None:
    """import_update_players does not overwrite an existing coach value."""
    conn = _make_db()
    # Pre-create a player with existing coach
    conn.execute(
        "INSERT INTO players (last_name, first_name, middle_name, birth_date, coach) "
        "VALUES (?, ?, ?, ?, ?)",
        ("иванов", "иван", None, "2010", "Существующий тренер"),
    )
    conn.commit()

    block = _make_block([
        {"fio": "Иванов Иван", "birth": "2010", "coach": "Новый тренер"},
    ])

    report = import_update_players(connection=conn, blocks=[block])

    assert report.unchanged == 1
    # Verify the coach was NOT changed
    row = conn.execute("SELECT coach FROM players WHERE last_name = 'иванов'").fetchone()
    assert row["coach"] == "Существующий тренер"


@pytest.mark.integration
def test_import_multi_tournament_creates_separate() -> None:
    """import_multi_tournament creates separate tournaments per block."""
    conn = _make_db()
    block1 = _make_block(
        [{"fio": "Иванов Иван", "birth": "2010", "place": "1", "coach": None}],
        sheet_name="Юниоры",
    )
    block2 = _make_block(
        [{"fio": "Петрова Мария", "birth": "2011", "place": "1", "coach": None}],
        sheet_name="Юниорки",
    )

    reports = import_multi_tournament(
        connection=conn,
        blocks=[block1, block2],
        base_name="Первенство",
        tournament_date="2025-01-01",
        is_adult_mode=False,
        source_files=["test.docx"],
    )

    assert len(reports) == 2
    # Verify tournaments have correct names
    tournaments = conn.execute(
        "SELECT name FROM tournaments ORDER BY id"
    ).fetchall()
    assert tournaments[0]["name"] == "Первенство - Юниоры"
    assert tournaments[1]["name"] == "Первенство - Юниорки"
